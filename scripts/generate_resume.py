#!/usr/bin/env python3
"""
generate_resume.py — Parse a resume markdown file and produce a polished
single-column .docx. Fully parametrized: NO personal data and NO letterhead
styling is hardcoded. Identity (name, contact, portfolio link) comes from the
resume markdown frontmatter and/or a profile.json; visual styling (fonts,
colors) comes from a format.json. Swap those two files and the same script
produces a different person's resume in a different visual identity.

WHY A SCRIPT: building the .docx by hand (python-docx calls, hairline rules,
tab-aligned dates, link styling) is fiddly and drifts every time it's re-derived.
This renders identically every run. The *content* is the model's job (tailor the
markdown); the *rendering* is this script's job.

Usage:
  python generate_resume.py --input resume_base.md --output resume.docx
  python generate_resume.py --input resume_acme.md --output Resume_Jane_Acme.docx \
      --profile ../user-library/profile.json --format ../user-library/format.json
  # optional company logo in the header:
  python generate_resume.py --input resume_acme.md --output out.docx \
      --logo logos/acme.png --logo-height 38

format.json (all keys optional — sensible defaults applied):
  {
    "name_font": "Canela Text",   # display font for the name
    "body_font": "Calibri",       # everything else
    "text_color": "1A1A1A",       # near-black: name, company, body
    "muted_color": "6B6B6B",      # gray: role titles, dates, contact
    "rule_color": "C8C8C8",       # hairline section rules / tab leader
    "accent_color": "175DC9"      # links (portfolio, projects)
  }

profile.json (only website_url is read here; the rest is used elsewhere):
  { "website_url": "https://example.com/" }

The resume markdown frontmatter carries the visible identity:
  ---
  name: Jane Doe
  title: [role-tailored title line]
  location: City, ST
  email: jane@example.com
  phone: (555) 123-4567
  website: example.com          # display text; link target is website_url
  ---
"""

import argparse
import json
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "ERROR: python-docx is not installed — the resume builder needs it.\n"
        "Install it with:  pip install python-docx\n"
        "(Run scripts/selfcheck.py to verify all dependencies at once.)\n"
    )
    sys.exit(3)

# ── Style tokens (set at runtime from format.json; these are the defaults) ──────
# Kept as module globals because the block builders below reference them. They
# are overwritten by apply_format() before any building happens.
TEXT  = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)
RULE  = RGBColor(0xC8, 0xC8, 0xC8)
LINK  = RGBColor(0x17, 0x5D, 0xC9)

FONT         = "Calibri"
NAME_FONT    = "Canela Text"
RULE_HEX     = "C8C8C8"
PORTFOLIO_URL = ""

CONTENT_W_IN  = 7.20
CONTENT_W_DXA = int(CONTENT_W_IN * 1440)


def _hex_to_rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def apply_format(fmt):
    """Overwrite the module style globals from a format dict (all keys optional)."""
    global TEXT, MUTED, RULE, LINK, FONT, NAME_FONT, RULE_HEX
    if fmt.get("text_color"):
        TEXT = _hex_to_rgb(fmt["text_color"])
    if fmt.get("muted_color"):
        MUTED = _hex_to_rgb(fmt["muted_color"])
    if fmt.get("rule_color"):
        RULE = _hex_to_rgb(fmt["rule_color"])
        RULE_HEX = fmt["rule_color"].lstrip("#").upper()
    if fmt.get("accent_color"):
        LINK = _hex_to_rgb(fmt["accent_color"])
    if fmt.get("body_font"):
        FONT = fmt["body_font"]
    if fmt.get("name_font"):
        NAME_FONT = fmt["name_font"]


def load_json(path):
    if path and os.path.exists(path):
        with open(path) as f:
            return json.load(f)
    return {}


# ── XML helpers ────────────────────────────────────────────────────────────────

def _add_right_tab_leader(paragraph, position_dxa=None, leader="none"):
    if position_dxa is None:
        position_dxa = CONTENT_W_DXA
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(qn('w:tabs')):
        pPr.remove(existing)
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(position_dxa)))
    tab.set(qn('w:leader'), leader)
    tabs.append(tab)
    pPr.append(tabs)


def _set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_dxa)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _set_table_width(table, width_dxa):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(width_dxa)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:val'), 'none')
        node.set(qn('w:sz'), '0')
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), 'auto')
        tblBorders.append(node)
    tblPr.append(tblBorders)


def _remove_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:val'), 'none')
        node.set(qn('w:sz'), '0')
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), 'auto')
        tcBorders.append(node)
    tcPr.append(tcBorders)


def _set_cell_margin(cell, top=0, start=0, bottom=0, end=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


# ── Style helpers ──────────────────────────────────────────────────────────────

def _run(p, text, size_pt, bold=False, italic=False, color=None, font=None):
    r = p.add_run(text)
    r.font.name = font or FONT
    r.font.size = Pt(size_pt)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def _fmt(p, space_before=0, space_after=0, line_spacing_pt=None,
         keep_with_next=False, left_indent_in=None, alignment=None):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if line_spacing_pt:
        pf.line_spacing = Pt(line_spacing_pt)
    if keep_with_next:
        pf.keep_with_next = True
    if left_indent_in is not None:
        pf.left_indent = Inches(left_indent_in)
    if alignment is not None:
        p.alignment = alignment


# ── Block builders ─────────────────────────────────────────────────────────────

def _add_top_border(paragraph, sz=4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(sz))
    top.set(qn('w:space'), '6')
    top.set(qn('w:color'), RULE_HEX)
    pBdr.append(top)
    pPr.append(pBdr)


def _section_header(doc, title):
    p = doc.add_paragraph()
    _fmt(p, space_before=14, space_after=3, keep_with_next=True)
    _add_top_border(p)
    _run(p, title.upper(), size_pt=8, bold=True, color=TEXT)
    return p


def _role_header(doc, company, role, date):
    p = doc.add_paragraph()
    _fmt(p, space_before=6, space_after=1, keep_with_next=True)
    _add_right_tab_leader(p)
    _run(p, company, size_pt=10.5, bold=True, color=TEXT)
    _run(p, '  /  ', size_pt=10.5, color=MUTED)
    _run(p, role, size_pt=10.5, color=MUTED)
    tab_r = p.add_run('\t')
    tab_r.font.name = FONT
    tab_r.font.size = Pt(10.5)
    if date:
        _run(p, date, size_pt=9, color=MUTED)
    return p


def _add_link(paragraph, text, url, size_pt=9, color_hex=None, underline=True):
    """Clickable hyperlink run, colored + optionally underlined."""
    r_id = paragraph.part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(sz)
    clr = OxmlElement('w:color')
    clr.set(qn('w:val'), color_hex or ('%02X%02X%02X' % (LINK[0], LINK[1], LINK[2])))
    rPr.append(clr)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    run_elem.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run_elem.append(t)
    hl.append(run_elem)
    paragraph._p.append(hl)


def _description(doc, text, space_after=3, is_last=False):
    p = doc.add_paragraph()
    _fmt(p, space_before=0, space_after=0 if is_last else space_after,
         line_spacing_pt=11, left_indent_in=0.18)
    _run(p, text, size_pt=9, color=TEXT)
    return p


# ── Markdown parser ────────────────────────────────────────────────────────────

def _parse_frontmatter(text):
    fm = {}
    if text.startswith('---'):
        end = text.index('---', 3)
        for line in text[3:end].strip().split('\n'):
            if ':' in line:
                k, v = line.split(':', 1)
                fm[k.strip()] = v.strip()
        text = text[end + 3:].strip()
    return fm, text


def _parse_sections(text):
    sections = {}
    sec_re = re.compile(r'^## (.+)$', re.MULTILINE)
    matches = list(sec_re.finditer(text))
    for i, m in enumerate(matches):
        name = m.group(1).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        sections[name] = text[start:end].strip()
    return sections


def _parse_experience(content):
    roles = []
    role_re = re.compile(r'^### (.+)$', re.MULTILINE)
    matches = list(role_re.finditer(content))
    for i, m in enumerate(matches):
        header = m.group(1).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
        body = content[start:end].strip()
        if ' | ' in header:
            role_line, date = header.rsplit(' | ', 1)
        else:
            role_line, date = header, ''
        if ' — ' in role_line:
            company, role = role_line.split(' — ', 1)
        else:
            company, role = role_line, ''
        lines = []
        for line in body.split('\n'):
            line = line.strip()
            if line.startswith('- '):
                lines.append(line[2:])
            elif line:
                lines.append(line)
        roles.append({'company': company.strip(), 'role': role.strip(),
                      'date': date.strip(), 'lines': lines})
    return roles


def _parse_projects(content):
    items = []
    for line in content.split('\n'):
        line = line.strip()
        if line.startswith('- '):
            url = ''
            inner = line[2:]
            if ' | http' in inner:
                inner, url = inner.rsplit(' | ', 1)
                url = url.strip()
            m = re.match(r'\*\*(.+?)\*\*\s*[—\-]+\s*(.*)', inner)
            if m:
                items.append({'name': m.group(1).strip(), 'desc': m.group(2).strip(), 'url': url})
            else:
                items.append({'name': inner, 'desc': '', 'url': url})
    return items


def _parse_skills(content):
    skills = []
    for line in content.split('\n'):
        m = re.match(r'\*\*(.+?):\*\*\s*(.*)', line.strip())
        if m:
            skills.append({'label': m.group(1).strip(), 'content': m.group(2).strip()})
    return skills


def _parse_education(content):
    return [l.strip()[2:] for l in content.split('\n') if l.strip().startswith('- ')]


def parse_resume(filepath):
    with open(filepath, 'r') as f:
        text = f.read()
    fm, body = _parse_frontmatter(text)
    sections = _parse_sections(body)
    return fm, sections


# ── Document assembler ─────────────────────────────────────────────────────────

def build_resume(fm, sections, profile=None, logo_path=None, logo_height_pt=40.0):
    profile = profile or {}
    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.top_margin    = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin   = Inches(0.65)
    sec.right_margin  = Inches(0.65)

    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after  = Pt(0)

    # Identity from frontmatter, falling back to profile.json.
    name     = fm.get('name', profile.get('name', ''))
    title    = fm.get('title', '')
    location = fm.get('location', profile.get('location', ''))
    email    = fm.get('email', profile.get('email', ''))
    phone    = fm.get('phone', profile.get('phone', ''))
    website  = fm.get('website', profile.get('website', ''))
    has_logo = logo_path and os.path.exists(logo_path)

    # Link target for the website display text: profile website_url, else https://<display>.
    portfolio_url = (PORTFOLIO_URL or profile.get('website_url')
                     or (('https://' + website) if website else ''))

    def _contact_paragraph(container):
        cp = container.add_paragraph() if hasattr(container, 'add_paragraph') else container
        _fmt(cp, space_before=0, space_after=0)
        prefix = "  ·  ".join(x for x in [location, email, phone] if x)
        _run(cp, prefix + ("  ·  " if website else ""), size_pt=8.5, color=MUTED)
        if website and portfolio_url:
            _add_link(cp, website, portfolio_url, size_pt=8.5)
        elif website:
            _run(cp, website, size_pt=8.5, color=MUTED)

    if has_logo:
        logo_col_dxa = 2376
        name_col_dxa = CONTENT_W_DXA - logo_col_dxa
        tbl = doc.add_table(rows=1, cols=2)
        _remove_table_borders(tbl)
        _set_table_width(tbl, CONTENT_W_DXA)
        left_cell  = tbl.cell(0, 0)
        right_cell = tbl.cell(0, 1)
        _set_cell_width(left_cell,  name_col_dxa)
        _set_cell_width(right_cell, logo_col_dxa)
        _remove_cell_borders(left_cell)
        _remove_cell_borders(right_cell)
        _set_cell_margin(left_cell,  top=0, start=0,  bottom=0, end=80)
        _set_cell_margin(right_cell, top=0, start=80, bottom=0, end=0)
        left_cell.vertical_alignment  = WD_ALIGN_VERTICAL.CENTER
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        np = left_cell.paragraphs[0]
        _fmt(np, space_before=0, space_after=2)
        _run(np, name, size_pt=20, color=TEXT, font=NAME_FONT)
        if title:
            tp = left_cell.add_paragraph()
            _fmt(tp, space_before=0, space_after=2)
            _run(tp, title, size_pt=9, color=MUTED)
        _contact_paragraph(left_cell)

        lp = right_cell.paragraphs[0]
        _fmt(lp, space_before=0, space_after=0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        lp.add_run().add_picture(logo_path, height=Pt(logo_height_pt))
    else:
        np = doc.add_paragraph()
        _fmt(np, space_before=0, space_after=2)
        _run(np, name, size_pt=20, color=TEXT, font=NAME_FONT)
        if title:
            tp = doc.add_paragraph()
            _fmt(tp, space_before=0, space_after=2)
            _run(tp, title, size_pt=9, color=MUTED)
        _contact_paragraph(doc)

    # Summary.
    summary_raw = sections.get('Summary', '').strip()
    if summary_raw:
        sp = doc.add_paragraph()
        _fmt(sp, space_before=8, space_after=0, line_spacing_pt=11)
        _run(sp, summary_raw, size_pt=9, color=TEXT)

    # Projects (inline).
    projects = _parse_projects(sections.get('Projects', ''))
    if projects:
        p = doc.add_paragraph()
        _fmt(p, space_before=8, space_after=3, line_spacing_pt=11)
        _run(p, 'Select Case Studies:  ', size_pt=9, bold=True, color=TEXT)
        for i, proj in enumerate(projects):
            if i > 0:
                _run(p, '  •  ', size_pt=9, color=MUTED)
            full_name = (proj['name'] + ' ' + proj['desc']).strip()
            url = proj.get('url', '')
            if url:
                _add_link(p, full_name, url, size_pt=9)
            else:
                _run(p, full_name, size_pt=9, color=TEXT)

    # Experience.
    if sections.get('Experience'):
        _section_header(doc, 'Professional Experience')
        for role in _parse_experience(sections.get('Experience', '')):
            _role_header(doc, role['company'], role['role'], role['date'])
            for i, line in enumerate(role['lines']):
                _description(doc, line, is_last=(i == len(role['lines']) - 1))

    # Skills.
    if sections.get('Skills'):
        _section_header(doc, 'Skills')
        for i, skill in enumerate(_parse_skills(sections.get('Skills', ''))):
            p = doc.add_paragraph()
            _fmt(p, space_before=3 if i > 0 else 1, space_after=0, line_spacing_pt=12)
            _run(p, skill['label'] + ':  ', size_pt=9, bold=True, color=TEXT)
            _run(p, skill['content'], size_pt=9, color=TEXT)

    # Education.
    if sections.get('Education'):
        _section_header(doc, 'Education')
        for i, item in enumerate(_parse_education(sections.get('Education', ''))):
            p = doc.add_paragraph()
            _fmt(p, space_before=3 if i > 0 else 1, space_after=0,
                 line_spacing_pt=12, left_indent_in=0.18)
            _run(p, item, size_pt=9, color=TEXT)

    return doc


# ── CLI ────────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    ap = argparse.ArgumentParser(description='Generate a resume .docx from markdown')
    ap.add_argument('--input',  required=True, help='Resume markdown file')
    ap.add_argument('--output', required=True, help='Output .docx path')
    ap.add_argument('--profile', default=None, help='Path to profile.json (contact/website_url defaults)')
    ap.add_argument('--format',  default=None, help='Path to format.json (fonts/colors)')
    ap.add_argument('--logo',    default=None, help='Path to a company logo (PNG/JPG)')
    ap.add_argument('--logo-height', type=float, default=40.0, help='Logo height in points')
    args = ap.parse_args()

    fmt = load_json(args.format)
    if fmt:
        apply_format(fmt)
    profile = load_json(args.profile)

    print(f'Parsing {args.input}...')
    fm, sections = parse_resume(args.input)
    print(f"Building resume for {fm.get('name', profile.get('name', '?'))}...")
    doc = build_resume(fm, sections, profile=profile,
                       logo_path=args.logo, logo_height_pt=args.logo_height)
    doc.save(args.output)
    print(f'Saved -> {args.output}')
