#!/usr/bin/env python3
"""
build_cover_letter.py — Render a cover letter to a polished .docx whose
letterhead matches the resume. NO personal data and NO styling is hardcoded:
the name font, body font, and accent color come from format.json (the same file
the resume uses), and the contact identity comes from the letter's frontmatter
(or profile.json). Swap those and the letterhead matches whoever's applying.

WHY THIS EXISTS: the cover letter and resume go in the same envelope, so they
must look like they came from the same person. Reusing the resume's design tokens
keeps them a matched set. Do NOT hand-write raw .docx or re-derive the styling.

The PROSE is not this script's job. The model writes the letter (in the user's
voice, per the skill) and saves it as markdown; this only renders that markdown.

INPUT markdown format:
    ---
    name: Jane Doe
    location: City, ST
    email: jane@example.com
    phone: (555) 123-4567
    website: example.com        # display text; link target is website_url below
    website_url: https://example.com/   # optional explicit link target
    date: June 2026             # optional
    salutation: Dear Hiring Team,   # optional
    signature: Jane Doe         # optional, defaults to name
    ---

    First paragraph...

    Second paragraph...

Usage:
    python build_cover_letter.py --input coverletter_<company>.md \
        --output "CoverLetter_<Name>_<Company>.docx" \
        --format ../user-library/format.json --profile ../user-library/profile.json
"""
import argparse
import json
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "ERROR: python-docx is not installed — the cover-letter builder needs it.\n"
        "Install it with:  pip install python-docx\n"
        "(Run scripts/selfcheck.py to verify all dependencies at once.)\n"
    )
    sys.exit(3)

# Defaults — overwritten from format.json so the letter matches the resume.
TEXT  = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)
LINK_HEX = "175DC9"
FONT = "Calibri"
NAME_FONT = "Canela Text"


def _hex_to_rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def load_json(path):
    if path and os.path.exists(path):
        with open(path) as f:
            return json.load(f)
    return {}


def apply_format(fmt):
    global TEXT, MUTED, LINK_HEX, FONT, NAME_FONT
    if fmt.get("text_color"):
        TEXT = _hex_to_rgb(fmt["text_color"])
    if fmt.get("muted_color"):
        MUTED = _hex_to_rgb(fmt["muted_color"])
    if fmt.get("accent_color"):
        LINK_HEX = fmt["accent_color"].lstrip("#").upper()
    if fmt.get("body_font"):
        FONT = fmt["body_font"]
    if fmt.get("name_font"):
        NAME_FONT = fmt["name_font"]


def _run(p, text, size_pt, bold=False, italic=False, color=None, font=None):
    r = p.add_run(text)
    r.font.name = font or FONT
    r.font.size = Pt(size_pt)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def _fmt(p, space_before=0, space_after=0, line_spacing_pt=None, alignment=None):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line_spacing_pt:
        pf.line_spacing = Pt(line_spacing_pt)
    if alignment is not None:
        p.alignment = alignment


def _add_link(paragraph, text, url, size_pt=9):
    """Accent-colored underlined hyperlink — same treatment as the resume."""
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), LINK_HEX)
    rPr.append(clr)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run_elem.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run_elem.append(t)
    hl.append(run_elem)
    paragraph._p.append(hl)


def parse(filepath):
    with open(filepath, "r") as f:
        text = f.read()
    fm, body = {}, text
    if text.startswith("---"):
        end = text.index("---", 3)
        for line in text[3:end].strip().split("\n"):
            if ":" in line:
                k, v = line.split(":", 1)
                fm[k.strip()] = v.strip()
        body = text[end + 3:].strip()
    paras = [p.strip().replace("\n", " ") for p in body.split("\n\n") if p.strip()]
    return fm, paras


def build(fm, paragraphs, profile=None):
    profile = profile or {}
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    def g(key):
        return fm.get(key, profile.get(key, ""))

    name = g("name")
    website = g("website")
    website_url = fm.get("website_url") or profile.get("website_url") or (
        ("https://" + website) if website else "")

    # Header — name in the display font, then the muted contact line.
    np = doc.add_paragraph()
    _fmt(np, space_before=0, space_after=2)
    _run(np, name, size_pt=20, color=TEXT, font=NAME_FONT)

    cp = doc.add_paragraph()
    _fmt(cp, space_before=0, space_after=0)
    contact = "  ·  ".join(x for x in [g("location"), g("email"), g("phone")] if x)
    _run(cp, contact + ("  ·  " if website else ""), size_pt=9, color=MUTED)
    if website and website_url:
        _add_link(cp, website, website_url, size_pt=9)
    elif website:
        _run(cp, website, size_pt=9, color=MUTED)

    if fm.get("date"):
        dp = doc.add_paragraph()
        _fmt(dp, space_before=14, space_after=0)
        _run(dp, fm["date"], size_pt=10.5, color=MUTED)

    if fm.get("salutation"):
        sp = doc.add_paragraph()
        _fmt(sp, space_before=14, space_after=0, line_spacing_pt=15)
        _run(sp, fm["salutation"], size_pt=10.5, color=TEXT)

    for i, para in enumerate(paragraphs):
        bp = doc.add_paragraph()
        _fmt(bp, space_before=10 if (i > 0 or fm.get("salutation")) else 14,
             space_after=0, line_spacing_pt=15)
        _run(bp, para, size_pt=10.5, color=TEXT)

    sig = doc.add_paragraph()
    _fmt(sig, space_before=14, space_after=0)
    _run(sig, fm.get("signature", name), size_pt=10.5, color=TEXT)
    return doc


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Render a cover letter .docx")
    ap.add_argument("--input", required=True)
    ap.add_argument("--output", required=True)
    ap.add_argument("--format", default=None, help="Path to format.json")
    ap.add_argument("--profile", default=None, help="Path to profile.json")
    args = ap.parse_args()

    fmt = load_json(args.format)
    if fmt:
        apply_format(fmt)
    profile = load_json(args.profile)

    fm, paragraphs = parse(args.input)
    if not paragraphs:
        print("ERROR: no body paragraphs found in input markdown", file=sys.stderr)
        sys.exit(1)
    doc = build(fm, paragraphs, profile=profile)
    doc.save(args.output)
    print(f"Saved -> {args.output}")
